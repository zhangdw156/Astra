# -*- coding: utf-8 -*-
"""
通用毕业论文格式修复脚本模板 v2.0
包含常见格式问题的修复函数集合，可按需组合使用
增强版：集成 DocumentStructure 位置感知，修复日志精确标注段落所在章节

用法：
  此脚本为模板，需根据具体学校规范要求自定义修复逻辑。
  Claude 在执行 Skill 工作流时会根据规范解析结果动态生成修复脚本。

  python fix_format.py <输入论文.docx> <输出论文.docx>
"""
import sys
import os

# 导入同目录下的 check_format 中的 DocumentStructure 和辅助函数
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from check_format import DocumentStructure, get_paragraph_indent_from_xml, get_spacing_from_xml

from docx import Document
from docx.shared import Pt, Cm, Emu, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

changes_log = []
_doc_structure = None  # 全局文档结构实例

# ★ v3.0新增：逐条修改记录，用于生成精细化对比报告
# 每条记录格式: {
#   'category': '页面设置',         # 修改大类
#   'item': '上边距',               # 检查项
#   'location': 'P38 [中文摘要]',   # 位置（段落号+章节）
#   'content_preview': '摘  要',    # 段落内容预览
#   'standard': '规范要求...',      # 规范要求
#   'before': '25.0mm',            # 修改前
#   'after': '28.0mm ✓'            # 修改后
# }
repair_records = []


def init_structure(doc):
    """初始化文档结构解析器（在修复开始前调用一次）"""
    global _doc_structure
    _doc_structure = DocumentStructure(doc)
    return _doc_structure


def get_para_location(para_index):
    """获取段落的位置描述，用于修复日志"""
    if _doc_structure is None:
        return ''
    loc = _doc_structure.get_location_short(para_index)
    return f' [{loc}]' if loc and loc != '?' else ''


def get_para_location_detail(para_index):
    """获取段落的详细位置信息，用于报告"""
    if _doc_structure is None:
        return f'P{para_index}', ''
    loc_short = _doc_structure.get_location_short(para_index)
    loc_full = _doc_structure.get_location(para_index)
    location_str = f'P{para_index} [{loc_short}]' if loc_short and loc_short != '?' else f'P{para_index}'
    return location_str, loc_full


def get_text_preview_short(doc, para_index, max_len=30):
    """获取段落的简短文本预览"""
    paras = doc.paragraphs
    if para_index >= len(paras):
        return ''
    text = paras[para_index].text.strip()
    if not text:
        return '(空段落)'
    if len(text) <= max_len:
        return text
    return text[:max_len] + '...'


def add_record(category, item, location, content_preview, standard, before, after):
    """添加一条逐条修改记录"""
    repair_records.append({
        'category': category,
        'item': item,
        'location': location,
        'content_preview': content_preview,
        'standard': standard,
        'before': before,
        'after': after,
    })


def log(msg):
    changes_log.append(msg)
    print(f"  {msg}")


# =====================================================
# 通用工具函数
# =====================================================

def ensure_ind(paragraph):
    """确保段落有 <w:ind> 元素"""
    pPr = paragraph._p.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = parse_xml(f'<w:ind {nsdecls("w")} />')
        pPr.append(ind)
    return ind


def set_first_line_zero(paragraph):
    """显式设首行缩进=0（覆盖样式继承）"""
    ind = ensure_ind(paragraph)
    for attr in ['firstLine', 'firstLineChars', 'hanging', 'hangingChars']:
        key = qn(f'w:{attr}')
        if ind.get(key) is not None:
            del ind.attrib[key]
    ind.set(qn('w:firstLine'), '0')
    ind.set(qn('w:firstLineChars'), '0')


def set_first_line_chars(paragraph, chars=2):
    """设首行缩进N字符"""
    ind = ensure_ind(paragraph)
    for attr in ['firstLine', 'firstLineChars', 'hanging', 'hangingChars']:
        key = qn(f'w:{attr}')
        if ind.get(key) is not None:
            del ind.attrib[key]
    ind.set(qn('w:firstLineChars'), str(int(chars * 100)))
    ind.set(qn('w:firstLine'), str(int(chars * 240)))


def set_hanging_indent_chars(paragraph, chars=2):
    """设悬挂缩进N字符"""
    ind = ensure_ind(paragraph)
    for attr in ['firstLine', 'firstLineChars', 'hanging', 'hangingChars']:
        key = qn(f'w:{attr}')
        if ind.get(key) is not None:
            del ind.attrib[key]
    ind.set(qn('w:hangingChars'), str(int(chars * 100)))
    ind.set(qn('w:hanging'), str(int(chars * 240)))
    ind.set(qn('w:leftChars'), str(int(chars * 100)))
    ind.set(qn('w:left'), str(int(chars * 240)))


def set_spacing(paragraph, before_lines=None, after_lines=None,
                before_pt=None, after_pt=None, line_spacing=None, line_rule=None):
    """设段前/段后间距（支持行数和磅值）"""
    pPr = paragraph._p.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = parse_xml(f'<w:spacing {nsdecls("w")} />')
        pPr.append(spacing)
    if before_lines is not None:
        spacing.set(qn('w:beforeLines'), str(int(before_lines * 100)))
        spacing.set(qn('w:before'), str(int(before_lines * 312)))
    if after_lines is not None:
        spacing.set(qn('w:afterLines'), str(int(after_lines * 100)))
        spacing.set(qn('w:after'), str(int(after_lines * 312)))
    if before_pt is not None:
        spacing.set(qn('w:before'), str(int(before_pt * 20)))
        # 清除 beforeLines 避免冲突
        if spacing.get(qn('w:beforeLines')) is not None:
            del spacing.attrib[qn('w:beforeLines')]
    if after_pt is not None:
        spacing.set(qn('w:after'), str(int(after_pt * 20)))
        if spacing.get(qn('w:afterLines')) is not None:
            del spacing.attrib[qn('w:afterLines')]
    if line_spacing is not None:
        if line_rule == 'exact':
            spacing.set(qn('w:line'), str(int(line_spacing * 20)))
            spacing.set(qn('w:lineRule'), 'exact')
        elif line_rule == 'atLeast':
            spacing.set(qn('w:line'), str(int(line_spacing * 20)))
            spacing.set(qn('w:lineRule'), 'atLeast')
        else:
            # 倍数行距
            spacing.set(qn('w:line'), str(int(line_spacing * 240)))
            spacing.set(qn('w:lineRule'), 'auto')


def remove_spacing(paragraph):
    """移除段前段后间距"""
    pPr = paragraph._p.find(qn('w:pPr'))
    if pPr is None:
        return
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        return
    for attr in ['before', 'beforeLines', 'after', 'afterLines']:
        key = qn(f'w:{attr}')
        if spacing.get(key) is not None:
            del spacing.attrib[key]


def set_run_font(run, cn_font='宋体', en_font='Times New Roman', size_pt=None, bold=None):
    """设置 run 的字体"""
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), en_font)
    rFonts.set(qn('w:hAnsi'), en_font)
    rFonts.set(qn('w:eastAsia'), cn_font)

    if size_pt is not None:
        half_pt = str(int(size_pt * 2))
        for tag in ['w:sz', 'w:szCs']:
            elem = rPr.find(qn(tag))
            if elem is None:
                elem = parse_xml(f'<{tag} {nsdecls("w")} />')
                rPr.append(elem)
            elem.set(qn('w:val'), half_pt)

    if bold is not None:
        run.font.bold = bold


def get_sectPr_list(doc):
    """获取所有节的 sectPr（含段落内嵌和body末尾的）"""
    sections = []
    body = doc.element.body
    for p in body.findall(qn('w:p')):
        pPr = p.find(qn('w:pPr'))
        if pPr is not None:
            sectPr = pPr.find(qn('w:sectPr'))
            if sectPr is not None:
                sections.append(sectPr)
    final = body.find(qn('w:sectPr'))
    if final is not None:
        sections.append(final)
    return sections


def set_line_spacing_for_range(doc, start_para, end_para, times=1.5, style_filter=None):
    """批量设置段落行距"""
    paras = doc.paragraphs
    count = 0
    location_groups = {}
    for i in range(start_para, min(end_para, len(paras))):
        p = paras[i]
        if style_filter and p.style.name != style_filter:
            continue
        if not p.text.strip():
            continue

        # 记录修改前
        old_spacing = get_spacing_from_xml(p)
        if 'line_spacing_times' in old_spacing:
            old_val = f'{old_spacing["line_spacing_times"]}倍'
        elif 'line_spacing_exact_pt' in old_spacing:
            old_val = f'固定{old_spacing["line_spacing_exact_pt"]}pt'
        elif 'line_spacing_atleast_pt' in old_spacing:
            old_val = f'最小{old_spacing["line_spacing_atleast_pt"]}pt'
        else:
            old_val = '继承'

        set_spacing(p, line_spacing=times)
        count += 1
        loc = get_para_location(i).strip(' []') or '未知'
        location_groups[loc] = location_groups.get(loc, 0) + 1

        loc_str, _ = get_para_location_detail(i)
        preview = get_text_preview_short(doc, i)
        style_info = f'(样式:{style_filter})' if style_filter else ''
        add_record('行距', f'行距{style_info}', loc_str, preview,
                   f'{times}倍行距', old_val, f'{times}倍 ✓')

    loc_detail = ', '.join(f'{k}({v}个)' for k, v in list(location_groups.items())[:8])
    log(f"设置了 {count} 个段落的行距为 {times} 倍 | 分布: {loc_detail}")


# =====================================================
# 修复函数集合
# =====================================================

def fix_page_margins(doc, top_mm=28, bottom_mm=22, left_mm=30, right_mm=20,
                     header_cm=1.8, footer_cm=1.4, gutter_mm=0,
                     section_indices=None):
    """
    修复页面边距
    section_indices: 要修复的节索引列表，None 表示全部
    """
    print("\n=== 修复页面边距 ===")
    sections = get_sectPr_list(doc)
    mm_to_twips = 1440 / 25.4

    for i, sect in enumerate(sections):
        if section_indices is not None and i not in section_indices:
            continue
        pgMar = sect.find(qn('w:pgMar'))
        if pgMar is None:
            continue

        # 记录修改前的值
        old_top = round(int(pgMar.get(qn('w:top'), '0')) / mm_to_twips, 1)
        old_bottom = round(int(pgMar.get(qn('w:bottom'), '0')) / mm_to_twips, 1)
        old_left = round(int(pgMar.get(qn('w:left'), '0')) / mm_to_twips, 1)
        old_right = round(int(pgMar.get(qn('w:right'), '0')) / mm_to_twips, 1)
        old_header = round(int(pgMar.get(qn('w:header'), '0')) / mm_to_twips / 10, 2)
        old_footer = round(int(pgMar.get(qn('w:footer'), '0')) / mm_to_twips / 10, 2)
        old_gutter = round(int(pgMar.get(qn('w:gutter'), '0')) / mm_to_twips, 1)

        pgMar.set(qn('w:top'), str(int(round(top_mm * mm_to_twips))))
        pgMar.set(qn('w:bottom'), str(int(round(bottom_mm * mm_to_twips))))
        pgMar.set(qn('w:left'), str(int(round(left_mm * mm_to_twips))))
        pgMar.set(qn('w:right'), str(int(round(right_mm * mm_to_twips))))
        pgMar.set(qn('w:header'), str(int(round(header_cm * 10 * mm_to_twips))))
        pgMar.set(qn('w:footer'), str(int(round(footer_cm * 10 * mm_to_twips))))
        pgMar.set(qn('w:gutter'), str(int(round(gutter_mm * mm_to_twips))))

        loc = f'节{i}'
        # 逐条记录每项变化
        margin_checks = [
            ('上边距', old_top, top_mm, 'mm'),
            ('下边距', old_bottom, bottom_mm, 'mm'),
            ('左边距', old_left, left_mm, 'mm'),
            ('右边距', old_right, right_mm, 'mm'),
            ('页眉距', old_header, header_cm, 'cm'),
            ('页脚距', old_footer, footer_cm, 'cm'),
            ('装订线', old_gutter, gutter_mm, 'mm'),
        ]
        for name, old_val, new_val, unit in margin_checks:
            if abs(old_val - new_val) > 0.05:
                add_record('页面设置', name, loc, f'节{i}页面设置',
                           f'{new_val}{unit}',
                           f'{old_val}{unit}',
                           f'{new_val}{unit} ✓')

        log(f"节{i}: 页边距 上{top_mm}/下{bottom_mm}/左{left_mm}/右{right_mm}mm, "
            f"页眉{header_cm}cm, 页脚{footer_cm}cm, 装订线{gutter_mm}mm")


def fix_paragraph_indent(doc, para_indices, chars=2, mode='first_line'):
    """
    修复段落缩进
    mode: 'first_line'=首行缩进, 'hanging'=悬挂缩进, 'zero'=无缩进
    """
    paras = doc.paragraphs
    count = 0
    location_groups = {}
    mode_names = {'first_line': '首行缩进', 'hanging': '悬挂缩进', 'zero': '无缩进'}
    mode_cn = mode_names.get(mode, mode)

    for idx in para_indices:
        if idx >= len(paras):
            continue
        p = paras[idx]

        # 记录修改前
        old_indent = get_paragraph_indent_from_xml(p)
        if mode == 'first_line':
            if 'first_line_chars' in old_indent:
                old_val = f'{old_indent["first_line_chars"]}字符'
            elif 'first_line_cm' in old_indent:
                old_val = f'{old_indent["first_line_cm"]}cm(固定值)'
            else:
                old_val = '无'
        elif mode == 'hanging':
            if 'hanging_chars' in old_indent:
                old_val = f'{old_indent["hanging_chars"]}字符'
            elif 'hanging_cm' in old_indent:
                old_val = f'{old_indent["hanging_cm"]}cm'
            else:
                old_val = '无'
        else:
            old_val = '有缩进' if old_indent else '无'

        if mode == 'first_line':
            set_first_line_chars(p, chars)
        elif mode == 'hanging':
            set_hanging_indent_chars(p, chars)
        elif mode == 'zero':
            set_first_line_zero(p)
        count += 1

        loc = get_para_location(idx).strip(' []') or '未知'
        location_groups[loc] = location_groups.get(loc, 0) + 1
        loc_str, _ = get_para_location_detail(idx)
        preview = get_text_preview_short(doc, idx)
        new_val = f'{chars}字符 ✓' if mode != 'zero' else '0(无缩进) ✓'

        add_record('段落缩进', mode_cn, loc_str, preview,
                   f'{chars}字符' if mode != 'zero' else '无缩进',
                   old_val, new_val)

    loc_detail = ', '.join(f'{k}({v}个)' for k, v in location_groups.items())
    log(f"修复了 {count} 个段落的缩进 (mode={mode}, chars={chars}) | 分布: {loc_detail}")


def fix_paragraph_spacing(doc, para_indices, before_lines=None, after_lines=None,
                          before_pt=None, after_pt=None):
    """修复段落间距"""
    paras = doc.paragraphs
    details = []
    for idx in para_indices:
        if idx >= len(paras):
            continue
        # 记录修改前
        old_spacing = get_spacing_from_xml(paras[idx])
        old_parts = []
        if 'before_lines' in old_spacing:
            old_parts.append(f'段前{old_spacing["before_lines"]}行')
        elif 'before_pt' in old_spacing:
            old_parts.append(f'段前{old_spacing["before_pt"]}pt')
        if 'after_lines' in old_spacing:
            old_parts.append(f'段后{old_spacing["after_lines"]}行')
        elif 'after_pt' in old_spacing:
            old_parts.append(f'段后{old_spacing["after_pt"]}pt')
        old_val = ', '.join(old_parts) if old_parts else '无'

        set_spacing(paras[idx], before_lines=before_lines, after_lines=after_lines,
                    before_pt=before_pt, after_pt=after_pt)

        new_parts = []
        if before_lines is not None: new_parts.append(f'段前{before_lines}行')
        if after_lines is not None: new_parts.append(f'段后{after_lines}行')
        if before_pt is not None: new_parts.append(f'段前{before_pt}pt')
        if after_pt is not None: new_parts.append(f'段后{after_pt}pt')
        new_val = ', '.join(new_parts) + ' ✓' if new_parts else '已设置 ✓'

        loc = get_para_location(idx)
        loc_str, _ = get_para_location_detail(idx)
        preview = get_text_preview_short(doc, idx)
        details.append(f"P{idx}{loc}")

        add_record('段落间距', '段前段后间距', loc_str, preview,
                   new_val.replace(' ✓', ''), old_val, new_val)

    log(f"修复了 {len(details)} 个段落的间距: {', '.join(details[:10])}" +
        (f' ...等{len(details)}个' if len(details) > 10 else ''))


def fix_paragraph_alignment(doc, para_indices, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    """修复段落对齐"""
    paras = doc.paragraphs
    details = []
    align_names = {0: 'LEFT', 1: 'CENTER', 2: 'RIGHT', 3: 'JUSTIFY', 4: 'DISTRIBUTE'}
    align_name = align_names.get(int(alignment), str(alignment))

    for idx in para_indices:
        if idx >= len(paras):
            continue
        old_align = paras[idx].alignment
        old_name = align_names.get(int(old_align) if old_align is not None else -1, '继承')
        paras[idx].alignment = alignment
        loc = get_para_location(idx)
        loc_str, _ = get_para_location_detail(idx)
        preview = get_text_preview_short(doc, idx)
        details.append(f"P{idx}{loc}")

        add_record('段落对齐', '对齐方式', loc_str, preview,
                   align_name, old_name, f'{align_name} ✓')

    log(f"修复了 {len(details)} 个段落对齐为 {align_name}: {', '.join(details[:10])}" +
        (f' ...等{len(details)}个' if len(details) > 10 else ''))


def fix_headers(doc, odd_text, even_text=None,
                cn_font='宋体', en_font='Times New Roman', size_pt=9,
                section_indices=None):
    """
    修复页眉
    odd_text: 奇数页页眉文字（通常是论文题目）
    even_text: 偶数页页眉文字（如"○○大学本科毕业论文"），None则不处理偶数页
    section_indices: 要修复的节索引列表，None 表示全部
    """
    print("\n=== 修复页眉 ===")
    fixed = 0
    fixed_sections = []
    from check_format import pt_to_chinese_size
    size_name = pt_to_chinese_size(size_pt)

    for i, section in enumerate(doc.sections):
        if section_indices is not None and i not in section_indices:
            continue
        try:
            header = section.header
            if header and header.paragraphs:
                for para in header.paragraphs:
                    if para.text.strip() or fixed == 0:
                        # 记录修改前
                        old_text = para.text.strip() or '(空)'
                        old_align_val = para.alignment
                        old_align = {0: '左对齐', 1: '居中', 2: '右对齐', 3: '两端对齐'}.get(
                            int(old_align_val) if old_align_val is not None else -1, '继承')
                        old_size = ''
                        if para.runs:
                            old_font = para.runs[0].font
                            if old_font.size:
                                old_size = pt_to_chinese_size(old_font.size.pt)

                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in para.runs:
                            run.text = ""
                        if para.runs:
                            para.runs[0].text = odd_text
                            set_run_font(para.runs[0], cn_font, en_font, size_pt)
                        elif odd_text:
                            run = para.add_run(odd_text)
                            set_run_font(run, cn_font, en_font, size_pt)
                        fixed += 1
                        fixed_sections.append(f"节{i}")

                        # 逐条记录
                        loc = f'节{i}页眉'
                        if old_text != odd_text:
                            add_record('页眉', '页眉内容', loc, f'页眉文字',
                                       f'完整论文题目', old_text[:40], f'{odd_text[:40]} ✓')
                        if old_align != '居中':
                            add_record('页眉', '页眉对齐', loc, f'页眉对齐方式',
                                       '居中对齐', old_align, '居中 ✓')
                        if old_size and old_size != size_name:
                            add_record('页眉', '页眉字号', loc, f'页眉字号',
                                       size_name, old_size, f'{size_name} ✓')
        except Exception as e:
            log(f"WARNING: 节{i}页眉修复失败: {e}")

        if even_text:
            try:
                even_header = section.even_page_header
                if even_header and even_header.paragraphs:
                    for para in even_header.paragraphs:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in para.runs:
                            run.text = ""
                        if para.runs:
                            para.runs[0].text = even_text
                            set_run_font(para.runs[0], cn_font, en_font, size_pt)
            except Exception as e:
                log(f"WARNING: 节{i}偶数页页眉修复失败: {e}")

    log(f"修复了 {fixed} 个节的页眉 ({', '.join(fixed_sections)}) → '{odd_text[:30]}...' 居中对齐")


def fix_tables_three_line(doc, top_bottom_sz=12, header_line_sz=8, table_indices=None):
    """
    将表格修复为三线表
    top_bottom_sz: 上下线粗细（1/8磅，12=1.5磅）
    header_line_sz: 标题行下线粗细（1/8磅，8=1磅）
    table_indices: 要修复的表格索引列表，None 表示全部
    """
    print("\n=== 修复表格为三线表 ===")
    
    # 构建表格位置映射
    table_locations = []
    body_elements = doc.element.body
    para_idx = 0
    for child in body_elements:
        if child.tag == qn('w:p'):
            para_idx += 1
        elif child.tag == qn('w:tbl'):
            table_locations.append(para_idx - 1 if para_idx > 0 else 0)
    
    for tbl_idx, table in enumerate(doc.tables):
        if table_indices is not None and tbl_idx not in table_indices:
            continue
        tbl = table._tbl
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            continue

        # 获取表格位置信息
        loc_info = ''
        if tbl_idx < len(table_locations):
            nearby_idx = table_locations[tbl_idx]
            loc_info = get_para_location(nearby_idx)
            # 尝试找表题
            for look_back in range(1, 6):
                check_idx = nearby_idx - look_back + 1
                if 0 <= check_idx < len(doc.paragraphs):
                    check_text = doc.paragraphs[check_idx].text.strip()
                    if check_text and ('表' in check_text and ('.' in check_text or '．' in check_text)):
                        loc_info += f" 表题:'{check_text[:40]}'"
                        break

        # 清除现有边框
        old = tblPr.find(qn('w:tblBorders'))
        if old is not None:
            tblPr.remove(old)

        # 三线表边框：上下粗线，无左右线，无内部线
        borders = parse_xml(f'<w:tblBorders {nsdecls("w")} />')
        borders.append(parse_xml(
            f'<w:top {nsdecls("w")} w:val="single" w:sz="{top_bottom_sz}" '
            f'w:space="0" w:color="auto"/>'))
        borders.append(parse_xml(
            f'<w:bottom {nsdecls("w")} w:val="single" w:sz="{top_bottom_sz}" '
            f'w:space="0" w:color="auto"/>'))
        borders.append(parse_xml(
            f'<w:left {nsdecls("w")} w:val="none" w:sz="0" w:space="0" w:color="auto"/>'))
        borders.append(parse_xml(
            f'<w:right {nsdecls("w")} w:val="none" w:sz="0" w:space="0" w:color="auto"/>'))
        borders.append(parse_xml(
            f'<w:insideH {nsdecls("w")} w:val="none" w:sz="0" w:space="0" w:color="auto"/>'))
        borders.append(parse_xml(
            f'<w:insideV {nsdecls("w")} w:val="none" w:sz="0" w:space="0" w:color="auto"/>'))
        tblPr.append(borders)

        # 标题行底部细线
        if len(table.rows) > 0:
            for cell in table.rows[0].cells:
                tcPr = cell._tc.get_or_add_tcPr()
                old_b = tcPr.find(qn('w:tcBorders'))
                if old_b is not None:
                    tcPr.remove(old_b)
                tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")} />')
                tcBorders.append(parse_xml(
                    f'<w:bottom {nsdecls("w")} w:val="single" w:sz="{header_line_sz}" '
                    f'w:space="0" w:color="auto"/>'))
                tcPr.append(tcBorders)

        log(f"表格{tbl_idx}{loc_info}: 三线表 (上下{top_bottom_sz/8}磅, 标题行下{header_line_sz/8}磅)")

        # 逐条记录
        add_record('表格', '表格样式→三线表', f'表格{tbl_idx}',
                   loc_info.strip() if loc_info else f'表格{tbl_idx}',
                   '三线表（上下粗线+标题行细线）',
                   '普通全边框',
                   f'三线表 ✓')


def fix_reference_title_spacing(doc, title_text="参  考  文  献"):
    """修复参考文献标题的字间距（插入全角空格）"""
    for i, para in enumerate(doc.paragraphs):
        cleaned = para.text.replace(' ', '').replace('\u3000', '')
        if cleaned == '参考文献' and (
            'Heading' in para.style.name
            or para.alignment == WD_ALIGN_PARAGRAPH.CENTER
        ):
            old_text = para.text.strip()
            if para.runs:
                for run in para.runs:
                    run.text = ""
                para.runs[0].text = title_text
                log(f"参考文献标题: -> '{title_text}'")
                loc_str, _ = get_para_location_detail(i)
                add_record('参考文献', '标题字间距', loc_str, old_text,
                           '字间加空格"参  考  文  献"',
                           f'"{old_text}"',
                           f'"{title_text}" ✓')
            break


def fix_reference_indent(doc, ref_title_keyword='参考文献', chars=2,
                         cn_font='宋体', en_font='Times New Roman', size_pt=None):
    """
    修复参考文献正文的悬挂缩进和字体
    ref_title_keyword: 用于定位参考文献标题的关键词（去空格后匹配）
    """
    ref_start = None
    paras = doc.paragraphs
    for i, p in enumerate(paras):
        cleaned = p.text.replace(' ', '').replace('\u3000', '')
        if cleaned == ref_title_keyword and 'Heading' in p.style.name:
            ref_start = i
            break

    if ref_start is None:
        log("WARNING: 未找到参考文献标题段落")
        return

    count = 0
    ref_entries = []
    for i in range(ref_start + 1, len(paras)):
        p = paras[i]
        if p.style.name.startswith('Heading'):
            break
        if not p.text.strip():
            continue

        # 记录修改前
        old_indent = get_paragraph_indent_from_xml(p)
        if 'hanging_chars' in old_indent:
            old_val = f'悬挂{old_indent["hanging_chars"]}字符'
        elif 'hanging_cm' in old_indent:
            old_val = f'悬挂{old_indent["hanging_cm"]}cm'
        elif 'first_line_chars' in old_indent:
            old_val = f'首行缩进{old_indent["first_line_chars"]}字符'
        elif 'first_line_cm' in old_indent:
            old_val = f'首行缩进{old_indent["first_line_cm"]}cm'
        else:
            old_val = '无缩进'

        set_hanging_indent_chars(p, chars=chars)
        for run in p.runs:
            set_run_font(run, cn_font, en_font, size_pt)
        count += 1
        ref_entries.append(f"P{i}")

        loc_str, _ = get_para_location_detail(i)
        preview = get_text_preview_short(doc, i, 35)
        add_record('参考文献', '悬挂缩进+字体', loc_str, preview,
                   f'悬挂缩进{chars}字符, 宋体/TNR',
                   old_val,
                   f'悬挂{chars}字符 ✓')

    loc = get_para_location(ref_start)
    log(f"修复了 {count} 条参考文献{loc}的悬挂缩进({chars}字符)和字体 "
        f"(P{ref_start+1}~P{ref_start+count})")


def fix_body_first_line_indent(doc, start_para, end_para, chars=2,
                               target_style='Normal', skip_centered=True):
    """
    批量修复正文段落首行缩进
    skip_centered: 是否跳过居中对齐的段落（如标题、图题）
    """
    paras = doc.paragraphs
    count = 0
    location_groups = {}
    for i in range(start_para, min(end_para, len(paras))):
        p = paras[i]
        if target_style and p.style.name != target_style:
            continue
        if not p.text.strip():
            continue
        if skip_centered and p.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            continue

        # 记录修改前
        old_indent = get_paragraph_indent_from_xml(p)
        if 'first_line_chars' in old_indent:
            old_val = f'{old_indent["first_line_chars"]}字符'
        elif 'first_line_cm' in old_indent:
            old_val = f'{old_indent["first_line_cm"]}cm(固定值)'
        else:
            old_val = '无'

        set_first_line_chars(p, chars)
        count += 1
        loc = get_para_location(i).strip(' []') or '未知'
        location_groups[loc] = location_groups.get(loc, 0) + 1

        loc_str, _ = get_para_location_detail(i)
        preview = get_text_preview_short(doc, i)
        add_record('正文首行缩进', '首行缩进', loc_str, preview,
                   f'{chars}字符', old_val, f'{chars}字符 ✓')

    loc_detail = ', '.join(f'{k}({v}个)' for k, v in location_groups.items())
    log(f"修复了 {count} 个正文段落的首行缩进为 {chars} 字符 | 分布: {loc_detail}")


def fix_body_extra_spacing(doc, start_para, end_para, target_style='Normal'):
    """移除正文段落多余的段前段后间距"""
    paras = doc.paragraphs
    count = 0
    affected = []
    for i in range(start_para, min(end_para, len(paras))):
        p = paras[i]
        if target_style and p.style.name != target_style:
            continue
        if not p.text.strip():
            continue

        pPr = p._p.find(qn('w:pPr'))
        if pPr is None:
            continue
        spacing = pPr.find(qn('w:spacing'))
        if spacing is None:
            continue

        bl = spacing.get(qn('w:beforeLines'))
        al = spacing.get(qn('w:afterLines'))
        b = spacing.get(qn('w:before'))
        a = spacing.get(qn('w:after'))

        old_parts = []
        if bl and bl != '0':
            old_parts.append(f'段前{int(bl)/100}行')
        if al and al != '0':
            old_parts.append(f'段后{int(al)/100}行')

        changed = False
        if bl and bl != '0':
            del spacing.attrib[qn('w:beforeLines')]
            if b:
                del spacing.attrib[qn('w:before')]
            changed = True
        if al and al != '0':
            del spacing.attrib[qn('w:afterLines')]
            if a:
                del spacing.attrib[qn('w:after')]
            changed = True
        if changed:
            count += 1
            loc = get_para_location(i)
            affected.append(f"P{i}{loc}")

            loc_str, _ = get_para_location_detail(i)
            preview = get_text_preview_short(doc, i)
            old_val = ', '.join(old_parts) if old_parts else '有多余间距'
            add_record('正文间距', '移除多余段间距', loc_str, preview,
                       '正文无段前段后', old_val, '已移除 ✓')

    if count:
        log(f"移除了 {count} 个正文段落的多余段间距: {', '.join(affected[:8])}" +
            (f' ...等{count}个' if count > 8 else ''))


def fix_font_for_range(doc, start_para, end_para,
                       cn_font='宋体', en_font='Times New Roman', size_pt=12,
                       bold=None, target_style=None):
    """批量修复段落字体"""
    paras = doc.paragraphs
    count = 0
    location_groups = {}
    for i in range(start_para, min(end_para, len(paras))):
        p = paras[i]
        if target_style and p.style.name != target_style:
            continue
        if not p.text.strip():
            continue
        for run in p.runs:
            set_run_font(run, cn_font, en_font, size_pt, bold)
        count += 1
        loc = get_para_location(i).strip(' []') or '未知'
        location_groups[loc] = location_groups.get(loc, 0) + 1
    
    loc_detail = ', '.join(f'{k}({v}个)' for k, v in list(location_groups.items())[:8])
    log(f"修复了 {count} 个段落的字体: 中文={cn_font}, 英文={en_font}, "
        f"字号={size_pt}pt, 加粗={bold} | 分布: {loc_detail}")


def fix_title_spacing(doc, para_idx, title_text=None, before_lines=None,
                      after_lines=None, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                      zero_indent=True):
    """
    修复标题段落（摘要、目录、致谢、参考文献等特殊标题）
    可设置文字内容、段前段后、对齐、缩进
    """
    paras = doc.paragraphs
    if para_idx >= len(paras):
        log(f"WARNING: 段落索引 {para_idx} 超出范围")
        return
    p = paras[para_idx]
    loc = get_para_location(para_idx)
    loc_str, _ = get_para_location_detail(para_idx)
    original_text = p.text.strip()[:40]

    if title_text is not None:
        if title_text != p.text.strip():
            add_record('特殊标题', '标题文字', loc_str, original_text,
                       f'"{title_text}"', f'"{original_text}"', f'"{title_text}" ✓')
        for run in p.runs:
            run.text = ""
        if p.runs:
            p.runs[0].text = title_text
        else:
            p.add_run(title_text)

    if alignment is not None:
        old_align = p.alignment
        align_names = {0: '左对齐', 1: '居中', 2: '右对齐', 3: '两端对齐'}
        old_name = align_names.get(int(old_align) if old_align is not None else -1, '继承')
        new_name = align_names.get(int(alignment), str(alignment))
        if old_name != new_name:
            add_record('特殊标题', '标题对齐', loc_str, original_text,
                       new_name, old_name, f'{new_name} ✓')
        p.alignment = alignment

    if zero_indent:
        old_indent = get_paragraph_indent_from_xml(p)
        had_indent = bool(old_indent.get('first_line_chars') or old_indent.get('first_line_cm'))
        set_first_line_zero(p)
        if had_indent:
            if 'first_line_chars' in old_indent:
                old_val = f'{old_indent["first_line_chars"]}字符'
            elif 'first_line_cm' in old_indent:
                old_val = f'{old_indent["first_line_cm"]}cm'
            else:
                old_val = '有'
            add_record('特殊标题', '标题首行缩进', loc_str, original_text,
                       '无缩进(0)', old_val, '0(无缩进) ✓')

    if before_lines is not None or after_lines is not None:
        old_spacing = get_spacing_from_xml(p)
        old_parts = []
        if 'before_lines' in old_spacing:
            old_parts.append(f'段前{old_spacing["before_lines"]}行')
        elif 'before_pt' in old_spacing:
            old_parts.append(f'段前{old_spacing["before_pt"]}pt')
        if 'after_lines' in old_spacing:
            old_parts.append(f'段后{old_spacing["after_lines"]}行')
        elif 'after_pt' in old_spacing:
            old_parts.append(f'段后{old_spacing["after_pt"]}pt')
        old_val = ', '.join(old_parts) if old_parts else '无'

        set_spacing(p, before_lines=before_lines, after_lines=after_lines)

        new_parts = []
        if before_lines is not None:
            new_parts.append(f'段前{before_lines}行')
        if after_lines is not None:
            new_parts.append(f'段后{after_lines}行')
        new_val = ', '.join(new_parts)
        add_record('特殊标题', '标题间距', loc_str, original_text,
                   new_val, old_val, f'{new_val} ✓')

    log(f"P{para_idx}{loc}: 标题修复 '{original_text}' → "
        f"(文字={title_text}, 段前={before_lines}行, 段后={after_lines}行)")


def fix_replace_text(doc, para_idx, old_text, new_text):
    """替换指定段落中的文本"""
    paras = doc.paragraphs
    if para_idx >= len(paras):
        return
    p = paras[para_idx]
    loc = get_para_location(para_idx)
    loc_str, _ = get_para_location_detail(para_idx)
    for run in p.runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)
            log(f"P{para_idx}{loc}: '{old_text}' → '{new_text}'")
            add_record('文本替换', '文本内容', loc_str, p.text.strip()[:30],
                       f'"{new_text}"', f'"{old_text}"', f'"{new_text}" ✓')


def fix_literature_type_label(doc, ref_title_keyword='参考文献',
                              replacements=None):
    """
    修复参考文献类型标识
    replacements: dict, 如 {'[硕士学位论文]': '[D]', '[期刊]': '[J]'}
    """
    if replacements is None:
        replacements = {
            '[硕士学位论文]': '[D]',
            '[博士学位论文]': '[D]',
            '[专著]': '[M]',
            '[期刊]': '[J]',
            '[会议论文]': '[C]',
            '[专利]': '[P]',
            '[报告]': '[R]',
            '[标准]': '[S]',
        }

    ref_start = None
    paras = doc.paragraphs
    for i, p in enumerate(paras):
        cleaned = p.text.replace(' ', '').replace('\u3000', '')
        if cleaned == ref_title_keyword:
            ref_start = i
            break

    if ref_start is None:
        return

    count = 0
    for i in range(ref_start + 1, len(paras)):
        p = paras[i]
        if p.style.name.startswith('Heading'):
            break
        for run in p.runs:
            for old, new in replacements.items():
                if old in run.text:
                    run.text = run.text.replace(old, new)
                    count += 1
                    loc_str, _ = get_para_location_detail(i)
                    preview = get_text_preview_short(doc, i, 35)
                    add_record('参考文献', '文献类型标识', loc_str, preview,
                               f'标准标识{new}', old, f'{new} ✓')
    if count:
        log(f"替换了 {count} 处文献类型标识")


# =====================================================
# 主函数模板
# =====================================================

def main():
    """
    这是一个模板主函数。
    实际使用时，Claude 会根据规范解析结果自动生成定制化的修复流程。
    """
    if len(sys.argv) < 3:
        print("用法: python fix_format.py <输入论文.docx> <输出论文.docx>")
        print("\n此脚本为模板(v2.0)，集成文档结构位置感知，包含以下可组合使用的修复函数：")
        print("  init_structure()               - 初始化文档结构解析（修复前必须调用）")
        print("  fix_page_margins()             - 修复页面边距")
        print("  fix_paragraph_indent()         - 修复段落缩进（含位置分布）")
        print("  fix_paragraph_spacing()        - 修复段落间距（含位置标注）")
        print("  fix_paragraph_alignment()      - 修复段落对齐（含位置标注）")
        print("  fix_headers()                  - 修复页眉（含节编号）")
        print("  fix_tables_three_line()        - 修复表格为三线表（含表格位置和表题）")
        print("  fix_reference_title_spacing()  - 修复参考文献标题字间距")
        print("  fix_reference_indent()         - 修复参考文献缩进（含段落范围）")
        print("  fix_body_first_line_indent()   - 批量修复正文首行缩进（含位置分布）")
        print("  fix_body_extra_spacing()       - 移除正文多余间距（含具体段落）")
        print("  fix_font_for_range()           - 批量修复字体（含位置分布）")
        print("  fix_title_spacing()            - 修复特殊标题段落（含位置标注）")
        print("  fix_replace_text()             - 替换段落文本（含位置标注）")
        print("  fix_literature_type_label()    - 修复文献类型标识")
        print("  set_run_font()                 - 设置单个 run 字体")
        print("  set_line_spacing_for_range()   - 批量设置行距（含位置分布）")
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2]

    print(f"Input:  {input_file}")
    print(f"Output: {output_file}")

    doc = Document(input_file)
    
    # ★ 初始化文档结构解析器（v2.0 必须步骤）
    structure = init_structure(doc)
    print(f"\n[文档结构] 共识别 {len(structure.structure)} 个章节标题/特殊节点")

    # ====== 以下为示例修复流程，需根据实际规范定制 ======
    #
    # 1. 页面设置
    # fix_page_margins(doc, top_mm=28, bottom_mm=22, left_mm=30, right_mm=20,
    #                  header_cm=1.8, footer_cm=1.4)
    #
    # 2. 页眉
    # fix_headers(doc, odd_text="论文完整题目（根据实际论文替换）")
    #
    # 3. 表格三线表
    # fix_tables_three_line(doc, top_bottom_sz=12, header_line_sz=8)
    #
    # 4. 参考文献
    # fix_reference_title_spacing(doc, title_text="参  考  文  献")
    # fix_reference_indent(doc, chars=2, cn_font='宋体', en_font='Times New Roman')
    # fix_literature_type_label(doc)
    #
    # 5. 正文首行缩进
    # fix_body_first_line_indent(doc, start_para=86, end_para=400, chars=2)
    #
    # 6. 移除正文多余间距
    # fix_body_extra_spacing(doc, start_para=86, end_para=400)
    #
    # 7. 摘要/致谢标题
    # fix_title_spacing(doc, para_idx=37, title_text=None,
    #                   before_lines=1, after_lines=1, zero_indent=True)
    #
    # ====================================================

    doc.save(output_file)
    print(f"\n[OK] Saved: {output_file}")
    print(f"\n{'=' * 60}")
    print(f" 修复摘要 ({len(changes_log)} 项)")
    print(f"{'=' * 60}")
    for i, c in enumerate(changes_log, 1):
        print(f"  {i:>2}. {c}")


if __name__ == '__main__':
    main()
