# -*- coding: utf-8 -*-
"""
论文格式修改前后对比报告生成器 v5.0
精简升级：
  - 只保留总结对比表格（删除原文+批注模式）
  - 新增"位置/定位"列，便于在原文中快速定位修改处
  - 表格 6 列：序号、位置/定位、检查项、规范文档要求、修改前、修改后

两种使用模式：
  模式1（推荐）: 从 fix_format.repair_records 自动生成
    from fix_format import repair_records
    create_report_from_records(output_path, title, subtitle, date_str, repair_records)

  模式2（兼容）: 手动传入 categories
    create_report(output_path, title, subtitle, date_str, total_changes, categories)
"""
import sys
import os
from collections import OrderedDict
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement


# =====================================================
# 表格格式工具函数
# =====================================================

def shade(cell, color):
    """设置单元格背景色"""
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>'))


def cell_fmt(cell, text, bold=False, sz=9, color=None, align=None,
             cn_font='宋体', en_font='Times New Roman'):
    """格式化单元格文字"""
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    # 设置段落间距紧凑
    pPr = p._p.get_or_add_pPr()
    sp = parse_xml(f'<w:spacing {nsdecls("w")} w:before="20" w:after="20" w:line="260" w:lineRule="auto"/>')
    pPr.append(sp)

    r = p.add_run(str(text))
    r.font.name = en_font
    r._element.rPr.rFonts.set(qn('w:eastAsia'), cn_font)
    r.font.size = Pt(sz)
    r.font.bold = bold
    if color:
        r.font.color.rgb = RGBColor(*color)


def hdr_row(tbl, cols, header_color="1F4E79"):
    """添加表头行"""
    row = tbl.add_row()
    for i, t in enumerate(cols):
        cell_fmt(row.cells[i], t, bold=True, sz=9, color=(255, 255, 255),
                 align=WD_ALIGN_PARAGRAPH.CENTER)
        shade(row.cells[i], header_color)
    return row


def cat_row(tbl, num, title, ncols=6, cat_color="D6E4F0"):
    """添加分类行（跨列合并）"""
    row = tbl.add_row()
    cell_fmt(row.cells[0], num, bold=True, sz=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    shade(row.cells[0], cat_color)
    for i in range(1, ncols):
        shade(row.cells[i], cat_color)
    # 合并 cells[1] ~ cells[ncols-1]
    for i in range(2, ncols):
        row.cells[1].merge(row.cells[i])
    cell_fmt(row.cells[1], title, bold=True, sz=9)
    return row


def data_row(tbl, no, location, item, standard, before, after):
    """添加数据行（6列：序号、位置/定位、检查项、规范文档要求、修改前、修改后）"""
    row = tbl.add_row()
    cell_fmt(row.cells[0], no, sz=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_fmt(row.cells[1], location, sz=8)
    cell_fmt(row.cells[2], item, sz=8)
    cell_fmt(row.cells[3], standard, sz=8)

    # 修改前 - 红色文字
    cell_fmt(row.cells[4], before, sz=8, color=(192, 0, 0))
    # 修改后
    if '✓' in str(after) or '✔' in str(after):
        cell_fmt(row.cells[5], after, sz=8, color=(0, 112, 0))
        shade(row.cells[5], "E2EFDA")
    elif '✗' in str(after) or '❌' in str(after):
        cell_fmt(row.cells[5], after, sz=8, color=(192, 0, 0))
        shade(row.cells[5], "FCE4EC")
    else:
        cell_fmt(row.cells[5], after, sz=8)

    return row


# =====================================================
# 记录分组与合并
# =====================================================

def _extract_location_info(locations):
    """
    从 location 列表中提取简洁的定位信息。
    返回 (简短位置文本, 内容预览文本)
    """
    if not locations:
        return '', ''

    # 提取段落号
    para_nums = []
    section_names = set()
    for loc in locations:
        if loc.startswith('P'):
            try:
                parts = loc.split(' ', 1)
                pnum = int(parts[0].replace('P', ''))
                para_nums.append(pnum)
                # 提取章节名 [xxx]
                if '[' in loc and ']' in loc:
                    sec = loc.split('[')[1].split(']')[0]
                    section_names.add(sec)
            except:
                pass
        elif loc.startswith('节'):
            section_names.add(loc)
        elif '页眉' in loc:
            section_names.add(loc)
        elif '表格' in loc:
            section_names.add(loc)
        else:
            section_names.add(loc)

    # 生成位置描述
    if para_nums:
        para_nums.sort()
        if len(para_nums) == 1:
            loc_text = f'第{para_nums[0]}段'
        elif len(para_nums) <= 3:
            loc_text = '、'.join(f'第{p}段' for p in para_nums)
        else:
            # 合并为范围
            loc_text = f'第{para_nums[0]}~{para_nums[-1]}段\n(共{len(para_nums)}处)'

        # 附带章节信息
        if section_names:
            secs = list(section_names)[:3]
            sec_str = '、'.join(secs)
            if len(section_names) > 3:
                sec_str += '等'
            loc_text += f'\n【{sec_str}】'
    elif section_names:
        loc_text = '、'.join(list(section_names)[:3])
    else:
        loc_text = '全局'

    return loc_text


def _extract_content_previews(records, max_previews=2, max_len=20):
    """从记录中提取内容预览文本"""
    previews = []
    seen = set()
    for rec in records:
        cp = rec.get('content_preview', '').strip()
        if cp and cp not in seen:
            seen.add(cp)
            if len(cp) > max_len:
                cp = cp[:max_len] + '…'
            previews.append(cp)
            if len(previews) >= max_previews:
                break
    return previews


def records_to_summary_categories(records):
    """
    将逐条修改记录按 category 分组，智能合并同类检查项，
    生成简洁的总结表格数据。

    合并规则：
    - 相同 category + item + standard + before + after 的记录合并
    - 合并后显示位置范围和内容预览，便于定位
    """
    # 按 category 分组
    cat_order = OrderedDict()
    for rec in records:
        cat_name = rec.get('category', '其他')
        if cat_name not in cat_order:
            cat_order[cat_name] = []
        cat_order[cat_name].append(rec)

    cn_nums = ['一', '二', '三', '四', '五', '六', '七', '八', '九', '十',
               '十一', '十二', '十三', '十四', '十五', '十六', '十七', '十八', '十九', '二十']

    categories = []
    global_no = 1

    for idx, (cat_name, items) in enumerate(cat_order.items()):
        cat = {
            'num': cn_nums[idx] if idx < len(cn_nums) else str(idx + 1),
            'title': cat_name,
            'items': []
        }

        # 合并同类型记录：按 (item, standard, before, after) 合并
        merged = OrderedDict()
        for rec in items:
            key = (rec.get('item', ''), rec.get('standard', ''),
                   rec.get('before', ''), rec.get('after', ''))
            if key not in merged:
                merged[key] = {
                    'item': rec.get('item', ''),
                    'standard': rec.get('standard', ''),
                    'before': rec.get('before', ''),
                    'after': rec.get('after', ''),
                    'locations': [],
                    'records': [],
                }
            merged[key]['locations'].append(rec.get('location', ''))
            merged[key]['records'].append(rec)

        for key, m in merged.items():
            loc_text = _extract_location_info(m['locations'])
            previews = _extract_content_previews(m['records'])

            # 如果有内容预览，附加在位置信息后
            if previews:
                preview_str = '\n'.join(f'"{p}"' for p in previews)
                if loc_text:
                    loc_text += f'\n{preview_str}'
                else:
                    loc_text = preview_str

            cat['items'].append({
                'no': str(global_no),
                'location': loc_text,
                'item': m['item'],
                'standard': m['standard'],
                'before': m['before'],
                'after': m['after'],
            })
            global_no += 1

        categories.append(cat)

    return categories


# =====================================================
# 核心报告构建
# =====================================================

def _build_summary_table(doc, categories, total_changes, date_str):
    """
    构建总结对比表格
    6列：序号、位置/定位、检查项、规范文档要求、修改前(原文件)、修改后(全面修改)
    """
    # 统计
    total_items = sum(len(cat['items']) for cat in categories)
    pass_items = sum(
        1 for cat in categories for item in cat['items']
        if '✓' in str(item['after'])
    )

    # 信息行
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'修改日期：{date_str}  ｜  共计 {total_changes} 项修改  ｜  '
                  f'合并为 {total_items} 项检查  ｜  已修复 {pass_items} 项')
    r.font.name = '宋体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(100, 100, 100)

    # 定位说明
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run('💡 提示：表中"位置/定位"列标注了段落编号（如"第88段"），'
                  '可在修复后的论文中通过 Ctrl+G 跳转至对应段落查看。')
    r.font.name = '宋体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(128, 128, 128)
    r.font.italic = True

    doc.add_paragraph()

    # 主表格（6列）
    tbl = doc.add_table(rows=0, cols=6)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    hdr_row(tbl, ['序号', '位置/定位', '检查项', '规范文档要求', '修改前（原文件）', '修改后（已修复）'])

    # 数据行
    row_counter = 0
    for cat in categories:
        cat_row(tbl, cat['num'], cat['title'])
        for item in cat['items']:
            r = data_row(tbl, item['no'], item.get('location', ''),
                         item['item'], item['standard'],
                         item['before'], item['after'])
            row_counter += 1
            # 交替行背景色
            if row_counter % 2 == 0:
                if '✓' not in item['after'] and '✗' not in item['after']:
                    for ci in range(6):
                        shade(r.cells[ci], "F5F5F5")

    # ========== 表格样式 ==========
    t = tbl._tbl
    tblPr = t.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")} />')
        t.insert(0, tblPr)

    # 边框
    old_b = tblPr.find(qn('w:tblBorders'))
    if old_b is not None:
        tblPr.remove(old_b)
    tblPr.append(parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="8" w:space="0" w:color="1F4E79"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="B4C6E7"/>'
        '<w:bottom w:val="single" w:sz="8" w:space="0" w:color="1F4E79"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="B4C6E7"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="D9E2F3"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="D9E2F3"/>'
        '</w:tblBorders>'))

    # 全宽
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblPr.append(parse_xml(
            f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>'))
    else:
        tblW.set(qn('w:w'), '5000')
        tblW.set(qn('w:type'), 'pct')

    # 列宽（6列）- A4 横向适配
    old_g = t.find(qn('w:tblGrid'))
    if old_g is not None:
        t.remove(old_g)
    t.insert(1, parse_xml(
        f'<w:tblGrid {nsdecls("w")}>'
        '<w:gridCol w:w="600"/>'    # 序号
        '<w:gridCol w:w="2200"/>'   # 位置/定位
        '<w:gridCol w:w="1800"/>'   # 检查项
        '<w:gridCol w:w="2800"/>'   # 规范文档要求
        '<w:gridCol w:w="2500"/>'   # 修改前
        '<w:gridCol w:w="2500"/>'   # 修改后
        '</w:tblGrid>'))


def _build_document(output_path, title, subtitle, date_str, total_changes,
                    categories, notes=None, basis=None):
    """
    构建 v5.0 对比报告文档（精简版：只有总结表格）
    """
    doc = Document()

    # ========== 页面设置（A4 横向，适配 6 列表格）==========
    for s in doc.sections:
        s.top_margin = Cm(2.0)
        s.bottom_margin = Cm(2.0)
        s.left_margin = Cm(1.8)
        s.right_margin = Cm(1.8)
        s.page_width = Cm(29.7)
        s.page_height = Cm(21.0)

    # ========== 标题 ==========
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.font.name = '黑体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    r.font.size = Pt(18)
    r.font.bold = True
    r.font.color.rgb = RGBColor(31, 78, 121)

    # ========== 副标题 ==========
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'《{subtitle}》')
    r.font.name = '宋体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(89, 89, 89)

    # ========== 总结对比表格 ==========
    _build_summary_table(doc, categories, total_changes, date_str)

    doc.add_paragraph()

    # ========== 分类统计摘要 ==========
    p = doc.add_paragraph()
    r = p.add_run('📊 修改统计')
    r.font.name = '黑体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = RGBColor(31, 78, 121)

    for cat in categories:
        total = len(cat['items'])
        passed = sum(1 for it in cat['items'] if '✓' in str(it['after']))
        p = doc.add_paragraph()
        if passed == total:
            marker = '✅'
            txt_color = RGBColor(0, 128, 0)
        else:
            marker = '⚠️'
            txt_color = RGBColor(192, 80, 77)
        r = p.add_run(f"  {marker} {cat['num']}、{cat['title']}: {passed}/{total} 项已修复")
        r.font.name = '宋体'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        r.font.size = Pt(10)
        r.font.color.rgb = txt_color

    doc.add_paragraph()

    # ========== 人工确认事项 ==========
    if notes:
        p = doc.add_paragraph()
        r = p.add_run('⚠️ 仍需人工确认的事项')
        r.font.name = '黑体'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        r.font.size = Pt(12)
        r.font.bold = True
        r.font.color.rgb = RGBColor(192, 80, 77)

        for n in notes:
            p = doc.add_paragraph()
            r = p.add_run(f'  • {n}')
            r.font.name = '宋体'
            r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            r.font.size = Pt(10)
            p._p.get_or_add_pPr().append(
                parse_xml(f'<w:ind {nsdecls("w")} w:leftChars="100" w:left="240"/>'))

        doc.add_paragraph()

    # ========== 修改依据 ==========
    if basis:
        p = doc.add_paragraph()
        r = p.add_run('📖 修改依据')
        r.font.name = '黑体'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        r.font.size = Pt(12)
        r.font.bold = True
        r.font.color.rgb = RGBColor(31, 78, 121)

        p = doc.add_paragraph()
        r = p.add_run(f'  {basis}')
        r.font.name = '宋体'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        r.font.size = Pt(10)
        p._p.get_or_add_pPr().append(
            parse_xml(f'<w:ind {nsdecls("w")} w:leftChars="100" w:left="240"/>'))

    # ========== 尾部统计 ==========
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    total_items = sum(len(cat['items']) for cat in categories)
    pass_items = sum(
        1 for cat in categories for item in cat['items']
        if '✓' in str(item['after']) or '✔' in str(item['after'])
    )
    fail_items = total_items - pass_items
    r = p.add_run(f'━━━  修改总计: {total_changes} 处  ┃  检查项: {total_items}  ┃  '
                  f'已修复: {pass_items}  ┃  待确认: {fail_items}  ━━━')
    r.font.name = '宋体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(89, 89, 89)

    doc.save(output_path)
    print(f'\n[OK] 对比报告已生成: {output_path}')
    print(f'     修改总计: {total_changes}, 检查项: {total_items}, '
          f'已修复: {pass_items}, 待确认: {fail_items}')
    return total_items, pass_items, fail_items


# =====================================================
# 外部接口
# =====================================================

def create_report_from_records(output_path, title, subtitle, date_str,
                                records, notes=None, basis=None,
                                doc_outline=None, fixed_doc_path=None):
    """
    v5.0 推荐接口：从 repair_records 列表自动生成对比报告

    Args:
        output_path: 输出文件路径
        title: 报告标题
        subtitle: 论文题目
        date_str: 日期字符串
        records: fix_format.repair_records 列表
        notes: 人工确认事项
        basis: 修改依据
        doc_outline: 文档结构概览（保留兼容，不使用）
        fixed_doc_path: 修复后论文路径（保留兼容，v5.0 不使用）
    """
    # 生成总结表格的 categories（合并同类项）
    categories = records_to_summary_categories(records)
    total_changes = len(records)
    return _build_document(output_path, title, subtitle, date_str,
                           total_changes, categories, notes, basis)


def create_report(output_path, title, subtitle, date_str, total_changes,
                  categories, notes=None, basis=None, doc_outline=None):
    """
    v2.0 兼容接口：手动传入 categories
    """
    return _build_document(output_path, title, subtitle, date_str,
                           total_changes, categories, notes, basis)


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("用法: python generate_report.py <输出路径.docx>")
        print("\n此脚本提供两个接口 (v5.0)：")
        print("  create_report_from_records() - ★推荐: 从 repair_records 自动生成")
        print("  create_report()              - 兼容v2.0: 手动传入 categories")
        sys.exit(0)

    # 示例数据（通用模板，不绑定特定院校）
    sample_records = [
        {'category': '页面设置', 'item': '上边距', 'location': '节0',
         'content_preview': '节0页面设置', 'standard': '上边距 28mm',
         'before': '25.0mm', 'after': '28.0mm ✓'},
        {'category': '页面设置', 'item': '下边距', 'location': '节0',
         'content_preview': '节0页面设置', 'standard': '下边距 22mm',
         'before': '25.0mm', 'after': '22.0mm ✓'},
        {'category': '中文摘要', 'item': '首行缩进', 'location': 'P39 [中文摘要]',
         'content_preview': '随着城市化进程的加速...', 'standard': '摘要正文段首缩进2字符',
         'before': '首行缩进 0.847cm', 'after': '首行缩进 2字符 ✓'},
    ]

    create_report_from_records(
        output_path=sys.argv[1],
        title='论文格式修改前后对比报告',
        subtitle='示例论文题目',
        date_str='2026年3月11日',
        records=sample_records,
        notes=['1. 请在Word中右键目录 → "更新域" → "更新整个目录"'],
        basis='《学校毕业论文（设计）撰写规范》（请替换为实际规范文档名称）',
    )
