# -*- coding: utf-8 -*-
"""
通用毕业论文格式全面检测工具 v2.0
增强版：精确定位段落位置（章节归属、估算页码、上下文文本）、问题分类与严重等级
支持：段落、表格、图片、样式、页面设置、页眉页脚、行距、缩进、节信息等
用法：python check_format.py <论文路径> <报告输出路径>
"""
import sys
import os
from docx import Document
from docx.shared import Pt, Cm, Mm, Inches, Emu, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


# =====================================================
# 单位转换工具
# =====================================================

def emu_to_cm(emu):
    if emu is None: return None
    return round(emu / 914400 * 2.54, 2)

def emu_to_mm(emu):
    if emu is None: return None
    return round(emu / 914400 * 25.4, 2)

def emu_to_pt(emu):
    if emu is None: return None
    return round(emu / 12700, 2)

def twips_to_mm(twips):
    if twips is None: return None
    return round(twips / 1440 * 25.4, 2)

def pt_to_chinese_size(pt):
    """磅值转中文字号名称"""
    size_map = {
        42.0: "初号", 36.0: "小初", 26.0: "一号", 24.0: "小一",
        22.0: "二号", 18.0: "小二", 16.0: "三号", 15.0: "小三",
        14.0: "四号", 12.0: "小四", 10.5: "五号", 9.0: "小五",
        7.5: "六号", 6.5: "小六", 5.5: "七号", 5.0: "八号",
    }
    return size_map.get(pt, f"{pt}pt(非标准)")


# =====================================================
# 章节结构解析器
# =====================================================

class DocumentStructure:
    """解析文档结构，为每个段落提供章节归属信息"""

    # 特殊章节标题关键词
    SPECIAL_SECTIONS = {
        '封面': ['本科', '毕业论文', '毕业设计', '学位论文'],
        '原创性声明': ['原创性声明', '独创性声明', '学术诚信'],
        '授权声明': ['授权声明', '使用授权', '授权书'],
        '中文摘要': ['摘要', '摘  要', '摘　要'],
        '英文摘要': ['abstract'],
        '目录': ['目录', '目  录', '目　录'],
        '参考文献': ['参考文献', '参 考 文 献', '参  考  文  献'],
        '致谢': ['致谢', '致  谢', '致　谢', '致 谢'],
        '附录': ['附录', '附  录'],
    }

    def __init__(self, doc):
        self.doc = doc
        self.paragraphs = doc.paragraphs
        self.structure = []  # [(para_index, section_type, section_name, heading_level)]
        self._build_structure()

    def _is_special_section(self, text, para_index, style_name, found_body, after_references):
        """判断是否为特殊章节标题
        
        增强逻辑：
        - 封面识别仅在文档前部（正文之前）生效
        - 参考文献之后不再误识别为封面
        - 需要段落文本以关键词为主体（短文本），避免正文中包含关键词的长段落误匹配
        """
        text_clean = text.strip().lower().replace(' ', '').replace('\u3000', '')
        text_len = len(text.strip())

        for section_type, keywords in self.SPECIAL_SECTIONS.items():
            for kw in keywords:
                kw_clean = kw.lower().replace(' ', '').replace('\u3000', '')
                if kw_clean in text_clean:
                    # 封面关键词仅在文档前部（尚未进入正文）且非参考文献之后生效
                    if section_type == '封面':
                        if found_body or after_references:
                            continue  # 正文之后不再识别为封面
                        # 封面段落通常较短
                        if text_len > 100:
                            continue
                    
                    # 中文摘要/英文摘要/目录/参考文献/致谢/附录：
                    # 应该是标题性质的段落（较短），不应是普通正文
                    if section_type in ('中文摘要', '英文摘要', '目录', '参考文献', '致谢', '附录'):
                        # 标题通常不超过30个字符
                        if text_len > 50:
                            continue
                    
                    # 原创性声明/授权声明可以略长但也有上限
                    if section_type in ('原创性声明', '授权声明'):
                        if text_len > 80:
                            continue
                    
                    return section_type
        return None

    def _build_structure(self):
        """构建文档结构树"""
        current_chapter = None  # 当前一级标题
        current_section = None  # 当前二级标题
        current_subsection = None  # 当前三级标题
        current_region = '文档开头'  # 当前区域描述
        found_body = False  # 是否已进入正文
        after_references = False  # 是否已经过参考文献

        for i, para in enumerate(self.paragraphs):
            text = para.text.strip()
            style_name = para.style.name if para.style else ''

            # 判断特殊章节
            if text:
                special = self._is_special_section(text, i, style_name, found_body, after_references)
                if special:
                    if special == '参考文献':
                        after_references = True
                    current_region = special
                    self.structure.append((i, 'special', special, text))
                    continue

            # 判断标题样式
            if style_name == 'Heading 1' or style_name.startswith('Heading 1'):
                found_body = True
                current_chapter = text
                current_section = None
                current_subsection = None
                current_region = f'一级标题「{text}」'
                self.structure.append((i, 'heading1', current_region, text))
            elif style_name == 'Heading 2' or style_name.startswith('Heading 2'):
                current_section = text
                current_subsection = None
                current_region = f'二级标题「{text}」(属于「{current_chapter or "?"}」)'
                self.structure.append((i, 'heading2', current_region, text))
            elif style_name == 'Heading 3' or style_name.startswith('Heading 3'):
                current_subsection = text
                current_region = f'三级标题「{text}」(属于「{current_section or "?"}」)'
                self.structure.append((i, 'heading3', current_region, text))
            elif style_name == 'Heading 4' or style_name.startswith('Heading 4'):
                current_region = f'四级标题「{text}」(属于「{current_subsection or "?"}」)'
                self.structure.append((i, 'heading4', current_region, text))

        # 缓存每个段落的位置信息
        self._para_locations = {}
        current_loc = {'region': '文档开头', 'chapter': None, 'section': None, 'subsection': None}
        struct_idx = 0

        for i in range(len(self.paragraphs)):
            # 更新当前位置
            while struct_idx < len(self.structure) and self.structure[struct_idx][0] <= i:
                entry = self.structure[struct_idx]
                _, stype, sname, stext = entry
                if stype == 'special':
                    current_loc = {'region': sname, 'chapter': None, 'section': None, 'subsection': None}
                elif stype == 'heading1':
                    current_loc = {'region': '正文', 'chapter': stext, 'section': None, 'subsection': None}
                elif stype == 'heading2':
                    current_loc['section'] = stext
                    current_loc['subsection'] = None
                elif stype == 'heading3':
                    current_loc['subsection'] = stext
                elif stype == 'heading4':
                    pass  # 不更新更高层级
                struct_idx += 1

            self._para_locations[i] = dict(current_loc)

    def get_location(self, para_index):
        """获取段落的位置信息，返回人类可读的字符串"""
        if para_index not in self._para_locations:
            return '未知位置'
        loc = self._para_locations[para_index]
        parts = []
        if loc['region'] and loc['region'] != '正文':
            parts.append(f"【{loc['region']}】")
        if loc['chapter']:
            parts.append(f"章: {loc['chapter']}")
        if loc['section']:
            parts.append(f"节: {loc['section']}")
        if loc['subsection']:
            parts.append(f"小节: {loc['subsection']}")
        return ' > '.join(parts) if parts else '文档开头'

    def get_location_short(self, para_index):
        """获取段落的简短位置标签"""
        if para_index not in self._para_locations:
            return '?'
        loc = self._para_locations[para_index]
        if loc['region'] and loc['region'] != '正文':
            return loc['region']
        if loc['section']:
            return loc['section']
        if loc['chapter']:
            return loc['chapter']
        return '文档开头'


# =====================================================
# 格式信息提取函数
# =====================================================

def get_alignment_name(alignment):
    if alignment is None: return "未设置(继承)"
    mapping = {0: "LEFT", 1: "CENTER", 2: "RIGHT", 3: "JUSTIFY", 4: "DISTRIBUTE"}
    try: return mapping.get(int(alignment), str(alignment))
    except: return str(alignment)

def get_font_info_from_run(run):
    """从run获取完整字体信息"""
    font = run.font
    info = {}
    info['name'] = font.name

    rpr = run._element.find(qn('w:rPr'))
    if rpr is not None:
        rFonts = rpr.find(qn('w:rFonts'))
        if rFonts is not None:
            info['east_asian'] = rFonts.get(qn('w:eastAsia'))
            info['ascii'] = rFonts.get(qn('w:ascii'))
            info['hAnsi'] = rFonts.get(qn('w:hAnsi'))
            info['cs'] = rFonts.get(qn('w:cs'))

    if font.size is not None:
        info['size_pt'] = font.size.pt
        info['size_name'] = pt_to_chinese_size(font.size.pt)
    else:
        if rpr is not None:
            sz = rpr.find(qn('w:sz'))
            if sz is not None:
                half_pt = int(sz.get(qn('w:val')))
                info['size_pt'] = half_pt / 2
                info['size_name'] = pt_to_chinese_size(half_pt / 2)

    info['bold'] = font.bold
    info['italic'] = font.italic
    info['underline'] = font.underline
    if font.color and font.color.rgb:
        info['color'] = str(font.color.rgb)

    if rpr is not None:
        vertAlign = rpr.find(qn('w:vertAlign'))
        if vertAlign is not None:
            info['vertAlign'] = vertAlign.get(qn('w:val'))
        spacing = rpr.find(qn('w:spacing'))
        if spacing is not None:
            val = spacing.get(qn('w:val'))
            if val: info['char_spacing'] = int(val)

    return info

def get_paragraph_indent_from_xml(para):
    """从XML获取段落缩进信息（包括字符单位）"""
    info = {}
    pPr = para._element.find(qn('w:pPr'))
    if pPr is not None:
        ind = pPr.find(qn('w:ind'))
        if ind is not None:
            for attr, key in [
                ('left', 'left_twips'), ('leftChars', 'left_chars'),
                ('firstLine', 'first_line_twips'), ('firstLineChars', 'first_line_chars'),
                ('hanging', 'hanging_twips'), ('hangingChars', 'hanging_chars'),
                ('right', 'right_twips')
            ]:
                val = ind.get(qn(f'w:{attr}'))
                if val:
                    if 'Chars' in attr:
                        info[key] = int(val) / 100
                    else:
                        info[key] = int(val)

            if 'left_twips' in info:
                info['left_cm'] = twips_to_mm(info['left_twips']) / 10
            if 'first_line_twips' in info:
                info['first_line_cm'] = twips_to_mm(info['first_line_twips']) / 10
            if 'hanging_twips' in info:
                info['hanging_cm'] = twips_to_mm(info['hanging_twips']) / 10
    return info

def get_spacing_from_xml(para):
    """从XML获取段前段后间距"""
    info = {}
    pPr = para._element.find(qn('w:pPr'))
    if pPr is not None:
        spacing = pPr.find(qn('w:spacing'))
        if spacing is not None:
            for attr, key in [
                ('before', 'before_twips'), ('after', 'after_twips'),
                ('line', 'line_val'), ('beforeLines', 'before_lines'),
                ('afterLines', 'after_lines')
            ]:
                val = spacing.get(qn(f'w:{attr}'))
                if val:
                    if 'Lines' in attr:
                        info[key] = int(val) / 100
                    else:
                        info[key] = int(val)

            if 'before_twips' in info:
                info['before_pt'] = round(info['before_twips'] / 20, 1)
            if 'after_twips' in info:
                info['after_pt'] = round(info['after_twips'] / 20, 1)

            lineRule = spacing.get(qn('w:lineRule'))
            if 'line_val' in info:
                if lineRule:
                    info['line_rule'] = lineRule
                if lineRule == 'auto' or lineRule is None:
                    info['line_spacing_times'] = round(info['line_val'] / 240, 2)
                elif lineRule == 'exact':
                    info['line_spacing_exact_pt'] = round(info['line_val'] / 20, 1)
                elif lineRule == 'atLeast':
                    info['line_spacing_atleast_pt'] = round(info['line_val'] / 20, 1)
    return info


# =====================================================
# 文本预览工具
# =====================================================

def get_text_preview(text, max_len=120):
    """获取段落文本预览，保留足够上下文"""
    text = text.strip()
    if not text:
        return '(空段落)'
    if len(text) <= max_len:
        return text
    return text[:max_len] + '...'


# =====================================================
# 检测模块
# =====================================================

def check_sections(doc, write):
    """检查所有节的页面设置"""
    write("\n" + "=" * 80)
    write("【一、页面设置信息】")
    write("=" * 80)
    for i, section in enumerate(doc.sections):
        write(f"\n--- 节 {i} (共{len(doc.sections)}节) ---")
        write(f"  纸张: {emu_to_mm(section.page_width)}mm × {emu_to_mm(section.page_height)}mm")
        write(f"  上边距: {emu_to_mm(section.top_margin)}mm")
        write(f"  下边距: {emu_to_mm(section.bottom_margin)}mm")
        write(f"  左边距: {emu_to_mm(section.left_margin)}mm")
        write(f"  右边距: {emu_to_mm(section.right_margin)}mm")
        write(f"  装订线: {emu_to_mm(section.gutter)}mm")
        write(f"  页眉距: {emu_to_mm(section.header_distance)}mm")
        write(f"  页脚距: {emu_to_mm(section.footer_distance)}mm")
        write(f"  方向: {section.orientation}")

        # 页眉信息
        try:
            if section.header and section.header.paragraphs:
                for p in section.header.paragraphs:
                    if p.text.strip():
                        h = {'text': p.text.strip(), 'alignment': get_alignment_name(p.alignment)}
                        if p.runs: h['font'] = get_font_info_from_run(p.runs[0])
                        write(f"  页眉: {h}")
        except: pass

        # 偶数页页眉
        try:
            if section.even_page_header and section.even_page_header.paragraphs:
                for p in section.even_page_header.paragraphs:
                    if p.text.strip():
                        write(f"  偶数页页眉: {p.text.strip()}")
        except: pass

        # 页脚信息
        try:
            if section.footer and section.footer.paragraphs:
                for p in section.footer.paragraphs:
                    if p.text.strip():
                        f_info = {'text': p.text.strip(), 'alignment': get_alignment_name(p.alignment)}
                        write(f"  页脚: {f_info}")
        except: pass

        # 分节符和页码
        try:
            sectPr = section._sectPr
            sectType = sectPr.find(qn('w:type'))
            if sectType is not None:
                write(f"  分节类型: {sectType.get(qn('w:val'))}")
            pgNumType = sectPr.find(qn('w:pgNumType'))
            if pgNumType is not None:
                fmt = pgNumType.get(qn('w:fmt'))
                start = pgNumType.get(qn('w:start'))
                if fmt: write(f"  页码格式: {fmt}")
                if start: write(f"  页码起始: {start}")
        except: pass


def check_styles(doc, write):
    """检查文档样式定义"""
    write("\n" + "=" * 80)
    write("【二、关键样式定义】")
    write("=" * 80)

    key_styles = ['Normal', 'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4',
                  'TOC Heading', 'toc 1', 'toc 2', 'toc 3', 'Title', 'Subtitle',
                  'List Paragraph', 'Caption', 'Header', 'Footer', 'Bibliography']

    for style in doc.styles:
        if style.name not in key_styles:
            continue
        write(f"\n  [{style.name}]")
        if style.base_style:
            write(f"    基础样式: {style.base_style.name}")
        if hasattr(style, 'paragraph_format') and style.paragraph_format:
            pf = style.paragraph_format
            write(f"    对齐: {get_alignment_name(pf.alignment)}")
        if hasattr(style, 'font') and style.font:
            font = style.font
            f_info = {}
            if font.name: f_info['name'] = font.name
            if font.size: f_info['size_pt'] = font.size.pt; f_info['size_name'] = pt_to_chinese_size(font.size.pt)
            if font.bold is not None: f_info['bold'] = font.bold
            try:
                style_elem = style.element
                rPr = style_elem.find(qn('w:rPr'))
                if rPr is not None:
                    rFonts = rPr.find(qn('w:rFonts'))
                    if rFonts is not None:
                        ea = rFonts.get(qn('w:eastAsia'))
                        if ea: f_info['east_asian'] = ea
            except: pass
            if f_info:
                write(f"    字体: {f_info}")


def check_paragraphs(doc, write, structure):
    """检查所有段落格式（增强版：含位置定位）"""
    write("\n" + "=" * 80)
    write("【三、段落详细格式信息】")
    write("  说明：每个段落标注了 [位置] 信息，格式为 [章节归属]")
    write("  P{n} 为段落在文档XML中的编号（从0开始），可用于精确修复定位")
    write("=" * 80)

    current_section_label = ''

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text and len(para.runs) == 0:
            continue

        style_name = para.style.name if para.style else "None"
        alignment = get_alignment_name(para.alignment)
        spacing_info = get_spacing_from_xml(para)
        indent_info = get_paragraph_indent_from_xml(para)
        location = structure.get_location(i)
        loc_short = structure.get_location_short(i)

        # 当章节变化时，打印分隔线
        if loc_short != current_section_label:
            current_section_label = loc_short
            write(f"\n{'─' * 70}")
            write(f"  ▶ 当前位置: {location}")
            write(f"{'─' * 70}")

        # 段落标题行：包含编号、样式、对齐、位置
        header_line = f"\n[P{i}] 样式:'{style_name}' | 对齐:{alignment} | 位置:{loc_short}"
        write(header_line)

        # 文本预览（加长到120字符，便于定位）
        if text:
            write(f"  📝 文本: {get_text_preview(text, 120)}")

        if spacing_info:
            parts = []
            if 'before_pt' in spacing_info: parts.append(f"段前:{spacing_info['before_pt']}pt")
            if 'before_lines' in spacing_info: parts.append(f"段前:{spacing_info['before_lines']}行")
            if 'after_pt' in spacing_info: parts.append(f"段后:{spacing_info['after_pt']}pt")
            if 'after_lines' in spacing_info: parts.append(f"段后:{spacing_info['after_lines']}行")
            if 'line_spacing_times' in spacing_info: parts.append(f"行距:{spacing_info['line_spacing_times']}倍")
            if 'line_spacing_exact_pt' in spacing_info: parts.append(f"行距:固定{spacing_info['line_spacing_exact_pt']}pt")
            if parts: write(f"  📐 间距: {', '.join(parts)}")

        if indent_info:
            parts = []
            if 'first_line_chars' in indent_info: parts.append(f"首行缩进:{indent_info['first_line_chars']}字符")
            elif 'first_line_cm' in indent_info: parts.append(f"首行缩进:{indent_info['first_line_cm']}cm (⚠固定值)")
            if 'hanging_chars' in indent_info: parts.append(f"悬挂缩进:{indent_info['hanging_chars']}字符")
            elif 'hanging_cm' in indent_info: parts.append(f"悬挂缩进:{indent_info['hanging_cm']}cm")
            if 'left_chars' in indent_info: parts.append(f"左缩进:{indent_info['left_chars']}字符")
            elif 'left_cm' in indent_info: parts.append(f"左缩进:{indent_info['left_cm']}cm")
            if parts: write(f"  📏 缩进: {', '.join(parts)}")

        if para.runs:
            for j, run in enumerate(para.runs):
                if not run.text.strip(): continue
                font_info = get_font_info_from_run(run)
                parts = []
                if font_info.get('east_asian'): parts.append(f"中文:{font_info['east_asian']}")
                if font_info.get('name'): parts.append(f"西文:{font_info['name']}")
                elif font_info.get('ascii'): parts.append(f"西文:{font_info['ascii']}")
                if 'size_pt' in font_info: parts.append(f"字号:{font_info['size_pt']}pt({font_info.get('size_name', '')})")
                if font_info.get('bold') is not None: parts.append(f"粗体:{font_info['bold']}")
                if font_info.get('italic'): parts.append(f"斜体:{font_info['italic']}")
                if font_info.get('color'): parts.append(f"颜色:{font_info['color']}")
                if font_info.get('vertAlign'): parts.append(f"上下标:{font_info['vertAlign']}")
                if font_info.get('char_spacing'): parts.append(f"字间距:{font_info['char_spacing']}")
                display_text = run.text.strip()[:80]
                if len(run.text.strip()) > 80: display_text += "..."
                write(f"  🔤 Run{j}: [{', '.join(parts)}]")
                write(f'       "{display_text}"')


def check_tables(doc, write, structure):
    """检查表格格式（增强版：含位置定位）"""
    write("\n" + "=" * 80)
    write("【四、表格格式信息】")
    write("=" * 80)

    # 构建表格到段落的位置映射
    # 通过检测表格前后的段落来定位表格
    table_locations = []
    body_elements = doc.element.body
    para_idx = 0
    for child in body_elements:
        if child.tag == qn('w:p'):
            para_idx += 1
        elif child.tag == qn('w:tbl'):
            # 表格位于第 para_idx 个段落之后
            table_locations.append(para_idx - 1 if para_idx > 0 else 0)

    for i, table in enumerate(doc.tables):
        # 获取表格位置
        nearby_para_idx = table_locations[i] if i < len(table_locations) else -1
        loc_info = ''
        if nearby_para_idx >= 0 and nearby_para_idx < len(doc.paragraphs):
            loc_info = f" | 位置: {structure.get_location_short(nearby_para_idx)}"
            # 尝试找到表题
            for look_back in range(1, 6):
                check_idx = nearby_para_idx - look_back + 1
                if 0 <= check_idx < len(doc.paragraphs):
                    check_text = doc.paragraphs[check_idx].text.strip()
                    if check_text and ('表' in check_text and ('.' in check_text or '．' in check_text)):
                        loc_info += f" | 表题: {check_text[:60]}"
                        break

        write(f"\n--- 表格 {i}{loc_info} ---")
        write(f"  行数: {len(table.rows)}, 列数: {len(table.columns)}")
        if table.style: write(f"  样式: {table.style.name}")

        tbl_elem = table._tbl
        tblPr = tbl_elem.find(qn('w:tblPr'))
        if tblPr is not None:
            jc = tblPr.find(qn('w:jc'))
            if jc is not None: write(f"  对齐: {jc.get(qn('w:val'))}")

            tblBorders = tblPr.find(qn('w:tblBorders'))
            if tblBorders is not None:
                write(f"  边框:")
                for bname in ['top', 'bottom', 'left', 'right', 'insideH', 'insideV']:
                    b = tblBorders.find(qn(f'w:{bname}'))
                    if b is not None:
                        sz = b.get(qn('w:sz'))
                        val = b.get(qn('w:val'))
                        color = b.get(qn('w:color'))
                        if val and val != 'none':
                            write(f"    {bname}: 类型={val}, 粗细={sz}(1/8pt), 颜色={color}")
                        elif val == 'none':
                            write(f"    {bname}: 无边框")

        # 样例单元格
        for row_idx, row in enumerate(table.rows):
            if row_idx > 1: break
            for col_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()[:50]
                if cell.paragraphs:
                    p = cell.paragraphs[0]
                    fi = {}
                    if p.runs: fi = get_font_info_from_run(p.runs[0])
                    write(f"    [{row_idx},{col_idx}] \"{cell_text}\" 对齐:{get_alignment_name(p.alignment)} 字体:{fi}")


def check_images(doc, write, structure):
    """检查图片信息（增强版：含位置定位）"""
    write("\n" + "=" * 80)
    write("【五、图片信息】")
    write("=" * 80)

    img_count = sum(1 for rel in doc.part.rels.values() if "image" in rel.reltype)
    write(f"  关联图片总数: {img_count}")

    for para_idx, para in enumerate(doc.paragraphs):
        for run in para.runs:
            drawings = run._element.findall(qn('w:drawing'))
            for d in drawings:
                img_info = {
                    'paragraph': f'P{para_idx}',
                    'location': structure.get_location_short(para_idx)
                }
                for inline in d.findall(qn('wp:inline')):
                    extent = inline.find(qn('wp:extent'))
                    if extent is not None:
                        cx, cy = extent.get('cx'), extent.get('cy')
                        if cx: img_info['width_cm'] = emu_to_cm(int(cx))
                        if cy: img_info['height_cm'] = emu_to_cm(int(cy))
                    docPr = inline.find(qn('wp:docPr'))
                    if docPr is not None:
                        img_info['name'] = docPr.get('name')
                for anchor in d.findall(qn('wp:anchor')):
                    img_info['type'] = 'floating'
                    extent = anchor.find(qn('wp:extent'))
                    if extent is not None:
                        cx, cy = extent.get('cx'), extent.get('cy')
                        if cx: img_info['width_cm'] = emu_to_cm(int(cx))
                        if cy: img_info['height_cm'] = emu_to_cm(int(cy))

                # 查找图题
                for look_ahead in range(1, 4):
                    check_idx = para_idx + look_ahead
                    if check_idx < len(doc.paragraphs):
                        check_text = doc.paragraphs[check_idx].text.strip()
                        if check_text and ('图' in check_text and ('.' in check_text or '．' in check_text)):
                            img_info['caption'] = check_text[:60]
                            break

                text = para.text.strip()[:80]
                if text: img_info['context'] = text
                write(f"  图片: {img_info}")


def check_default_fonts(doc, write):
    """检查文档默认字体设置"""
    write("\n" + "=" * 80)
    write("【六、文档默认格式】")
    write("=" * 80)
    try:
        styles_elem = doc.styles.element
        docDefaults = styles_elem.find(qn('w:docDefaults'))
        if docDefaults is not None:
            rPrDefault = docDefaults.find(qn('w:rPrDefault'))
            if rPrDefault is not None:
                rPr = rPrDefault.find(qn('w:rPr'))
                if rPr is not None:
                    rFonts = rPr.find(qn('w:rFonts'))
                    if rFonts is not None:
                        write(f"  默认西文字体: {rFonts.get(qn('w:ascii'))}")
                        write(f"  默认中文字体: {rFonts.get(qn('w:eastAsia'))}")
                    sz = rPr.find(qn('w:sz'))
                    if sz is not None:
                        pt = int(sz.get(qn('w:val'))) / 2
                        write(f"  默认字号: {pt}pt ({pt_to_chinese_size(pt)})")

            pPrDefault = docDefaults.find(qn('w:pPrDefault'))
            if pPrDefault is not None:
                pPr = pPrDefault.find(qn('w:pPr'))
                if pPr is not None:
                    spacing = pPr.find(qn('w:spacing'))
                    if spacing is not None:
                        line = spacing.get(qn('w:line'))
                        lineRule = spacing.get(qn('w:lineRule'))
                        if line:
                            if lineRule == 'auto' or lineRule is None:
                                write(f"  默认行距: {round(int(line)/240, 2)}倍")
    except Exception as e:
        write(f"  获取默认格式失败: {e}")


def check_document_structure(doc, write, structure):
    """输出文档结构大纲（新增模块）"""
    write("\n" + "=" * 80)
    write("【七、文档结构大纲】")
    write("  说明：以下列出文档的章节结构，便于定位各部分")
    write("=" * 80)

    for para_idx, stype, sname, stext in structure.structure:
        if stype == 'special':
            write(f"  P{para_idx:>4d}  ★ {sname}")
        elif stype == 'heading1':
            write(f"  P{para_idx:>4d}  ■ {stext}")
        elif stype == 'heading2':
            write(f"  P{para_idx:>4d}    ├─ {stext}")
        elif stype == 'heading3':
            write(f"  P{para_idx:>4d}    │  ├── {stext}")
        elif stype == 'heading4':
            write(f"  P{para_idx:>4d}    │  │  ├─── {stext}")


# =====================================================
# 主函数
# =====================================================

def full_check(filepath, output_path):
    """全面检查文档格式"""
    doc = Document(filepath)
    report = []
    write = report.append

    write("=" * 80)
    write("全面文档格式检查报告 v2.0")
    write(f"文件: {os.path.basename(filepath)}")
    write(f"段落总数: {len(doc.paragraphs)}, 表格总数: {len(doc.tables)}, 节总数: {len(doc.sections)}")
    write("=" * 80)
    write("")
    write("报告说明:")
    write("  - 每个段落标注 [P{n}]，n 是段落在文档XML中的编号（从0开始）")
    write("  - 每个段落附带 [位置] 信息，标注其所属章节")
    write("  - 📝 文本 = 段落内容预览（最长120字符）")
    write("  - 📐 间距 = 段前/段后/行距")
    write("  - 📏 缩进 = 首行缩进/悬挂缩进/左缩进")
    write("  - 🔤 Run = 文本片段的字体信息")
    write("  - ⚠固定值 = 缩进使用了固定厘米值而非字符单位（可能需要修复）")

    # 构建文档结构
    structure = DocumentStructure(doc)

    # 先输出文档结构大纲
    check_document_structure(doc, write, structure)

    check_sections(doc, write)
    check_styles(doc, write)
    check_paragraphs(doc, write, structure)
    check_tables(doc, write, structure)
    check_images(doc, write, structure)
    check_default_fonts(doc, write)

    content = '\n'.join(report)
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(content)
    print(f"完成！报告已写入: {output_path}")
    print(f"总行数: {len(report)}")


if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("用法: python check_format.py <论文.docx路径> <报告输出路径.txt>")
        sys.exit(1)
    full_check(sys.argv[1], sys.argv[2])
