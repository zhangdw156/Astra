import os
import sys
import zipfile
import argparse
from docx import Document

def get_image_format(part):
    content_type = part.content_type
    if 'jpeg' in content_type or 'jpg' in content_type:
        return 'jpg'
    elif 'gif' in content_type:
        return 'gif'
    elif 'bmp' in content_type:
        return 'bmp'
    return 'png'

def docx_to_md(doc_path, output_dir=None):
    if not os.path.exists(doc_path):
        print(f"错误: 文件不存在: {doc_path}")
        return False
    
    if output_dir is None:
        base_name = os.path.splitext(os.path.basename(doc_path))[0]
        output_dir = os.path.join(os.path.dirname(doc_path), f"{base_name}_output")
    
    os.makedirs(output_dir, exist_ok=True)
    
    image_files = []
    with zipfile.ZipFile(doc_path, 'r') as zip_ref:
        media_files = [f for f in zip_ref.namelist() if f.startswith('word/media/') and not f.endswith('/')]
        for i, img_file in enumerate(media_files):
            ext = os.path.splitext(img_file)[1]
            img_name = f"image_{i+1}{ext}"
            zip_ref.extract(img_file, output_dir)
            old_path = os.path.join(output_dir, img_file)
            new_path = os.path.join(output_dir, img_name)
            if os.path.exists(old_path):
                os.rename(old_path, new_path)
                image_files.append(img_name)
    
    try:
        doc = Document(doc_path)
    except Exception as e:
        print(f"错误: 无法读取Word文档: {e}")
        return False
    
    md_content = []
    doc_title = os.path.splitext(os.path.basename(doc_path))[0]
    md_content.append(f"# {doc_title}\n\n")
    
    def find_drawings_in_paragraph(para):
        drawings_found = []
        for run in para.runs:
            el = run._element
            drawings = el.xpath('.//w:drawing')
            if drawings:
                drawings_found.extend(drawings)
        return drawings_found
    
    image_idx = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        
        if not text:
            md_content.append('\n')
            continue
        
        style_name = para.style.name if para.style else ''
        
        if 'Heading 1' in style_name or style_name == '标题 1':
            md_content.append(f'\n# {text}\n')
        elif 'Heading 2' in style_name or style_name == '标题 2':
            md_content.append(f'\n## {text}\n')
        elif 'Heading 3' in style_name or style_name == '标题 3':
            md_content.append(f'\n### {text}\n')
        elif 'Heading 4' in style_name or style_name == '标题 4':
            md_content.append(f'\n#### {text}\n')
        elif style_name in ['Title', '标题']:
            md_content.append(f'\n# {text}\n')
        elif style_name in ['List Bullet', '列表']:
            md_content.append(f'- {text}\n')
        elif style_name in ['List Number', '编号列表']:
            md_content.append(f'1. {text}\n')
        else:
            md_content.append(f'{text}\n')
        
        drawings = find_drawings_in_paragraph(para)
        if drawings:
            for _ in drawings:
                if image_idx < len(image_files):
                    img_name = image_files[image_idx]
                    md_content.append(f'![{img_name}]({img_name})\n')
                    image_idx += 1
    
    for table in doc.tables:
        md_content.append('\n')
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip().replace('\n', ' ')
                row_data.append(cell_text)
            md_content.append('| ' + ' | '.join(row_data) + ' |\n')
        md_content.append('\n')
    
    if image_idx < len(image_files):
        md_content.append('\n\n## 附录：图片\n\n')
        for i in range(image_idx, len(image_files)):
            img_name = image_files[i]
            md_content.append(f'![{img_name}]({img_name})\n')
    
    md_path = os.path.join(output_dir, f'{doc_title}.md')
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write(''.join(md_content))
    
    print(f"转换完成!")
    print(f"MD文件: {md_path}")
    print(f"图片数量: {len(image_files)}")
    print(f"输出目录: {output_dir}")
    return True

def main():
    parser = argparse.ArgumentParser(description='将Word文档(.docx)转换为Markdown格式并提取图片')
    parser.add_argument('input_file', help='输入的Word文档路径')
    parser.add_argument('-o', '--output', help='输出目录 (可选)')
    
    args = parser.parse_args()
    
    success = docx_to_md(args.input_file, args.output)
    sys.exit(0 if success else 1)

if __name__ == '__main__':
    main()
